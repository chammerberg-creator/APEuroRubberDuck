"""
Rubber Duck — Feynman Practice (MVP)  — Streamlit Cloud friendly
[...trimmed in this comment: full content inserted again below...]
"""
import os
import re
import uuid
from datetime import datetime
from typing import List, Dict
from io import BytesIO

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt

APP_TITLE = "🦆 Rubber Duck — Feynman Practice (MVP)"
DEFAULT_MODEL = os.environ.get("OPENAI_CHAT_MODEL", "gpt-4o-mini")
VOCAB_CAP = 16

_api_key = (st.secrets.get("OPENAI_API_KEY", None) 
            if hasattr(st, "secrets") else None) or os.environ.get("OPENAI_API_KEY", None)
try:
    client = OpenAI(api_key=_api_key) if _api_key else OpenAI()
except Exception as e:
    client = None

def ensure_state():
    st.session_state.setdefault("active_scope_title", "")
    st.session_state.setdefault("chapters", [])
    st.session_state.setdefault("chat_history", [])
    st.session_state.setdefault("student_name", "")
    st.session_state.setdefault("chapter_limit", 1)
    st.session_state.setdefault("captured_terms", set())

def extract_candidate_terms(text: str):
    import re
    cleaned = re.sub(r"[^A-Za-z0-9\-\s]", " ", text)
    tokens = cleaned.split()
    candidates = []
    for i, tok in enumerate(tokens):
        if (tok.istitle() and i + 1 < len(tokens) and tokens[i + 1].istitle()) or ("-" in tok and len(tok) > 3):
            candidates.append(tok)
        elif tok.istitle() and len(tok) > 3:
            candidates.append(tok)
    seen = set(); out = []
    for t in candidates:
        if t.lower() not in seen:
            seen.add(t.lower()); out.append(t)
    return out[:VOCAB_CAP]

def build_system_prompt(scope_title: str, chapter_limit: int) -> str:
    return f"""
You are **Rubber Duck**, a curious peer using the Feynman Technique in a history class.
Ask, don't tell. One short question at a time. Stay strictly within the student's declared progress.

SCOPE_TITLE: {scope_title or '(not set)'}
STUDENT_PROGRESS: up to Chapter {chapter_limit}

Priorities (in order):
1) Missing definitions (ask for plain-language definitions of in-scope vocabulary)
2) Concrete examples (real people/events, data, places) — avoid abstractions
3) Cause → effect and contrasts (how is this similar/different, why?)
4) Significance (why it matters)

Rules:
- Never lecture the correct answer; guide with questions.
- Never ask about future chapters; if invited, politely refocus on the current scope.
- Occasionally (low frequency) restate something slightly wrong to invite correction.
- Do not quote textbooks or external sources. Use any provided excerpts only to choose better questions.
- Keep responses concise and friendly. Exactly one question per turn.
""".strip()

def next_rubber_duck_reply(scope_title: str, chapter_limit: int, history, model: str) -> str:
    if client is None:
        return "(OpenAI client is not configured. Add OPENAI_API_KEY in Streamlit Secrets or environment.)"
    system_prompt = build_system_prompt(scope_title, chapter_limit)
    hidden_context = "RETRIEVED_EXCERPTS (for internal guidance only; do not reveal):\n(none in MVP)"
    messages = [{"role": "system", "content": system_prompt},
                {"role": "system", "content": hidden_context},
                *history[-20:]]
    resp = client.chat.completions.create(
        model=model, messages=messages, temperature=0.4, max_tokens=220
    )
    return resp.choices[0].message.content.strip()

def render_docx_buffer(student_name, scope_title, chapter_limit, transcript, vocab_terms):
    from docx import Document
    from docx.shared import Pt
    from io import BytesIO
    from datetime import datetime

    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("Personalized Study Guide — Rubber Duck (Feynman)")
    run.bold = True; run.font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.add_run("Student: ").bold = True; meta.add_run(student_name or "[Name]")
    meta.add_run("    Date: ").bold = True; meta.add_run(datetime.now().strftime("%Y-%m-%d"))
    meta.add_run("    Scope: ").bold = True; meta.add_run(f"{scope_title or '[Chapter]'} — up to Chapter {chapter_limit}")

    doc.add_heading("1) Snapshot Summary", level=2)
    student_lines = [m["content"] for m in transcript if m["role"] == "user"][-4:]
    if student_lines:
        doc.add_paragraph("Here’s what you explained in this session:")
        for s in student_lines:
            doc.add_paragraph(f"• {s[:200]}" + ("…" if len(s) > 200 else ""))
    else:
        doc.add_paragraph("[Your summary here]")

    doc.add_heading("2) Strengths — Covered Well", level=2)
    doc.add_paragraph("• Provided specific examples when prompted.")
    doc.add_paragraph("• Clarified at least one definition in plain language.")

    doc.add_heading("3) Focus Areas — Needs Review", level=2)
    doc.add_paragraph("• Define key terms succinctly before details.")
    doc.add_paragraph("• Add cause → effect links and significance.")

    doc.add_heading("4) Vocabulary Status", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Term"; hdr[1].text = "Student’s Definition (own words)"; hdr[2].text = "Status (Mastered / Developing / Review)"
    for term in list(vocab_terms)[:VOCAB_CAP]:
        row = table.add_row().cells
        row[0].text = term; row[1].text = "[your definition]"; row[2].text = "Developing"

    doc.add_heading("8) Action Plan (Next Steps)", level=2)
    doc.add_paragraph("1) Re‑explain the core idea in one sentence.")
    doc.add_paragraph("2) Add one concrete example with who/when/where.")
    doc.add_paragraph("3) Write a cause → effect chain and one sentence on why it matters.")

    buf = BytesIO(); doc.save(buf); buf.seek(0); return buf

# UI
st.set_page_config(page_title="Rubber Duck — Feynman MVP", page_icon="🦆", layout="wide")
st.title(APP_TITLE); ensure_state()

with st.sidebar:
    st.header("Mode")
    mode = st.radio("Select mode", ["Teacher", "Student"], horizontal=True)
    model = st.text_input("Model (optional)", value=DEFAULT_MODEL, help="e.g., gpt-4o-mini")
    if not _api_key:
        st.warning("Add OPENAI_API_KEY in Streamlit Secrets (or set env var) to enable responses.")

if mode == "Teacher":
    st.subheader("Teacher: Scope & Chapters")
    scope_title = st.text_input("Scope title (e.g., 'Expansion of Europe, Units 1–4')", value=st.session_state.active_scope_title)
    st.write("**Upload chapter files (optional in MVP):** Their names become the chapter list. Sorting is by filename.")
    uploads = st.file_uploader("Chapters (PDF/TXT)", type=["pdf", "txt"], accept_multiple_files=True)
    if st.button("Activate Scope"):
        st.session_state.active_scope_title = scope_title
        chapters = []
        if uploads:
            for i, uf in enumerate(sorted(uploads, key=lambda x: x.name), start=1):
                chapters.append({"order": i, "id": uf.name})
        else:
            chapters = [{"order": i, "id": f"Chapter {i}"} for i in range(1, 6)]
        st.session_state.chapters = chapters
        st.success(f"Activated scope with {len(chapters)} chapter(s).")
    if st.session_state.chapters:
        st.markdown("**Chapters:**")
        st.write([f"{c['order']:02d} — {c['id']}" for c in st.session_state.chapters])
else:
    st.subheader("Student: Explain your chapter — I’m Rubber Duck 🦆")
    if not st.session_state.chapters:
        st.info("No active scope yet. Ask your teacher to activate chapters in Teacher mode."); st.stop()
    st.session_state.student_name = st.text_input("Your name", value=st.session_state.student_name)
    labels = [f"{c['order']:02d} — {c['id']}" for c in st.session_state.chapters]
    default_idx = min(len(labels) - 1, 0)
    choice = st.selectbox("Last chapter you finished", labels, index=default_idx)
    chapter_limit = int(choice.split(" — ")[0]) if choice else 1
    st.session_state.chapter_limit = chapter_limit
    st.caption("Starter line (optional): 'We’re practicing the Feynman Technique. I read [chapter] up to this point…'")

    for m in st.session_state.chat_history:
        st.chat_message("user" if m["role"] == "user" else "assistant").markdown(m["content"])

    user_msg = st.chat_input("Start explaining here… (I’ll ask one short question at a time)")
    if user_msg:
        st.session_state.chat_history.append({"role": "user", "content": user_msg})
        for term in extract_candidate_terms(user_msg):
            st.session_state.captured_terms.add(term)
        with st.chat_message("assistant"):
            try:
                reply = next_rubber_duck_reply(st.session_state.active_scope_title, chapter_limit, st.session_state.chat_history, model)
            except Exception as e:
                reply = f"(Error generating response: {e})"
            st.markdown(reply)
        st.session_state.chat_history.append({"role": "assistant", "content": reply})

    st.divider()
    if st.button("All done — generate my study guide"):
        vocab_terms = list(st.session_state.captured_terms)[:VOCAB_CAP]
        buf = render_docx_buffer(
            st.session_state.student_name,
            st.session_state.active_scope_title,
            st.session_state.chapter_limit,
            st.session_state.chat_history,
            vocab_terms,
        )
        filename = f"{(st.session_state.student_name or 'student').replace(' ', '_')}_rubber_duck_study_guide_{datetime.now():%Y%m%d_%H%M%S}.docx"
        st.download_button("Download study guide (.docx)", data=buf.getvalue(), file_name=filename,
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.success("Study guide generated! You can download it above.")
    st.caption("I’ll only ask about chapters you’ve finished. If I miss something, nudge me!")
